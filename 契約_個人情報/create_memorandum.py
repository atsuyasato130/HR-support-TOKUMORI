"""
覚書（人材紹介手数料変更）生成スクリプト
基本契約: 2025年10月9日付「人材紹介（新卒）に関する基本契約書」
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_cell_border(cell, **kwargs):
    """テーブルセルの罫線を設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = OxmlElement(tag)
            element.set(qn("w:val"), kwargs[edge].get("val", "single"))
            element.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), kwargs[edge].get("color", "000000"))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_background(cell, color_hex):
    """セルの背景色設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_paragraph_with_style(doc, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT,
                              space_before=0, space_after=0, indent=0, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "MS 明朝"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "MS 明朝")
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def create_memorandum():
    doc = Document()

    # ページ余白設定
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # デフォルトフォント設定
    style = doc.styles["Normal"]
    style.font.name = "MS 明朝"
    style.font.size = Pt(10.5)

    # ========== タイトル ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("覚　書")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "MS 明朝"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "MS 明朝")

    # ========== 前文 ==========
    preamble_text = (
        "　株式会社エール（以下「甲」という）と株式会社TOKUMORI（以下「乙」という）とは、"
        "2025年10月9日付にて締結した「人材紹介（新卒）に関する基本契約書」（以下「基本契約」という）"
        "第7条第3項の規定に基づき、2027年度新卒採用に係る紹介手数料について、以下のとおり覚書"
        "（以下「本覚書」という）を締結する。"
    )
    p_pre = doc.add_paragraph()
    p_pre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_pre.paragraph_format.space_before = Pt(0)
    p_pre.paragraph_format.space_after = Pt(12)
    run = p_pre.add_run(preamble_text)
    run.font.size = Pt(10.5)
    run.font.name = "MS 明朝"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "MS 明朝")

    # ========== 第1条 ==========
    add_paragraph_with_style(doc, "第１条（目的）", bold=True, size=10.5, space_before=6, space_after=4)
    add_paragraph_with_style(
        doc,
        "　本覚書は、基本契約第7条に定める紹介手数料について、2027年度新卒採用の採用区分ごとに"
        "個別の手数料額および適用条件を定めることを目的とする。",
        size=10.5, space_after=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 第2条 ==========
    add_paragraph_with_style(doc, "第２条（定義）", bold=True, size=10.5, space_before=6, space_after=4)
    add_paragraph_with_style(
        doc,
        "　本覚書において使用する用語の定義は以下のとおりとする。",
        size=10.5, space_after=4
    )
    definitions = [
        ("①", "「関東新卒採用」とは、採用後の配属・勤務地が関東地方（東京都・神奈川県・埼玉県・"
               "千葉県・茨城県・栃木県・群馬県）のいずれかの拠点となる2027年度新卒採用をいう。"),
        ("②", "「関西新卒採用」とは、採用後の配属・勤務地が関西地方（大阪府・兵庫県・京都府・"
               "奈良県・滋賀県・和歌山県）のいずれかの拠点となる2027年度新卒採用をいう。"),
        ("③", "「内定承諾」とは、基本契約第6条に定める「採用決定者が内定を承諾または業務委託契約"
               "締結に同意した段階」をいう。"),
    ]
    for num, text in definitions:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        run = p.add_run(f"{num}　{text}")
        run.font.size = Pt(10.5)
        run.font.name = "MS 明朝"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "MS 明朝")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ========== 第3条 ==========
    add_paragraph_with_style(doc, "第３条（紹介手数料および適用条件）", bold=True, size=10.5,
                              space_before=6, space_after=4)
    add_paragraph_with_style(
        doc,
        "　甲は、乙の人材紹介を通じて採用が決定した場合、採用区分に応じて以下の紹介手数料を"
        "乙に対して支払うものとする。なお、消費税は別途加算するものとする。",
        size=10.5, space_after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 手数料テーブル ==========
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"

    # 列幅設定（合計13cm程度）
    col_widths = [Cm(3.8), Cm(3.5), Cm(6.2)]
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # ヘッダー行
    header_row = table.rows[0]
    headers = ["採用区分", "紹介手数料（税別）", "適用条件"]
    for i, (cell, header) in enumerate(zip(header_row.cells, headers)):
        set_cell_background(cell, "C0C0C0")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = "MS 明朝"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "MS 明朝")

    # データ行1: 関東新卒採用
    row1 = table.rows[1]
    row1_data = [
        "2027年度\n関東新卒採用",
        "金 120万円/人",
        "・乙による紹介日が2026年4月1日以降であること\n"
        "・応募者の内定承諾日が2026年9月30日までであること\n"
        "・2027年3月卒業見込みで2027年4月1日入社予定の者",
    ]
    for i, (cell, text) in enumerate(zip(row1.cells, row1_data)):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 2 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.name = "MS 明朝"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "MS 明朝")

    # データ行2: 関西新卒採用
    row2 = table.rows[2]
    row2_data = [
        "2027年度\n関西新卒採用",
        "金 100万円/人",
        "・2027年3月卒業見込みで2027年4月1日入社予定の者\n"
        "・乙による紹介日および内定承諾日については、別途甲乙協議のうえ定める",
    ]
    for i, (cell, text) in enumerate(zip(row2.cells, row2_data)):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 2 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.name = "MS 明朝"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "MS 明朝")

    # 注記行
    row3 = table.rows[3]
    row3.cells[0].merge(row3.cells[2])
    note_cell = row3.cells[0]
    set_cell_background(note_cell, "F5F5F5")
    p = note_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(
        "※ 上記に定めのない採用区分・採用決定者については、基本契約第7条第1項（100万円/人）が適用される。"
    )
    run.font.size = Pt(9)
    run.font.name = "MS 明朝"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "MS 明朝")
    run.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ========== 第4条 ==========
    add_paragraph_with_style(doc, "第４条（支払方法）", bold=True, size=10.5, space_before=10, space_after=4)
    add_paragraph_with_style(
        doc,
        "　本覚書に定める紹介手数料の支払方法については、基本契約第8条の規定を準用する。",
        size=10.5, space_after=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 第5条 ==========
    add_paragraph_with_style(doc, "第５条（手数料の返還）", bold=True, size=10.5, space_before=6, space_after=4)
    add_paragraph_with_style(
        doc,
        "　本覚書に定める紹介手数料の返還については、基本契約第9条の規定を準用する。",
        size=10.5, space_after=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 第6条 ==========
    add_paragraph_with_style(doc, "第６条（内定報告義務）", bold=True, size=10.5, space_before=6, space_after=4)
    add_paragraph_with_style(
        doc,
        "　甲は、本覚書の適用対象となる採用決定者が内定を承諾した場合、基本契約第10条に定める"
        "内定報告義務を負う。報告義務の懈怠があった場合の取扱いについても、同条の規定に準ずる。",
        size=10.5, space_after=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 第7条 ==========
    add_paragraph_with_style(doc, "第７条（直接取引の禁止）", bold=True, size=10.5, space_before=6, space_after=4)
    add_paragraph_with_style(
        doc,
        "　本覚書の適用対象となる応募者についても、基本契約第11条（直接取引の禁止）の規定が"
        "適用される。",
        size=10.5, space_after=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 第8条 ==========
    add_paragraph_with_style(doc, "第８条（本覚書の優先）", bold=True, size=10.5, space_before=6, space_after=4)
    add_paragraph_with_style(
        doc,
        "　本覚書と基本契約に内容の相違がある場合は、本覚書の定めを優先する。本覚書に定めのない"
        "事項については、基本契約の規定が適用される。",
        size=10.5, space_after=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 第9条 ==========
    add_paragraph_with_style(doc, "第９条（有効期間）", bold=True, size=10.5, space_before=6, space_after=4)
    add_paragraph_with_style(
        doc,
        "　本覚書の有効期間は、本覚書の締結日から、本覚書に定める全ての採用決定者に係る紹介手数料"
        "の支払いおよび精算が完了するまでの期間とする。",
        size=10.5, space_after=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 第10条 ==========
    add_paragraph_with_style(doc, "第１０条（協議）", bold=True, size=10.5, space_before=6, space_after=4)
    add_paragraph_with_style(
        doc,
        "　本覚書に定めのない事項または解釈に疑義が生じた場合は、基本契約第21条に従い、甲及び乙"
        "が誠実に協議のうえ解決するものとする。",
        size=10.5, space_after=16, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 締結文 ==========
    add_paragraph_with_style(
        doc,
        "　本覚書の成立を証するため、本書の電磁的記録を作成し、甲及び乙が合意の後電子署名を施し、"
        "各自その電磁的記録を保管する。",
        size=10.5, space_after=20, align=WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    # ========== 日付 ==========
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.space_after = Pt(16)
    run = date_p.add_run("　　　　　　年　　月　　日")
    run.font.size = Pt(11)
    run.font.name = "MS 明朝"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "MS 明朝")

    # ========== 署名欄 ==========
    # 甲
    add_paragraph_with_style(doc, "甲", bold=True, size=11, space_before=4, space_after=3)
    for line in [
        "　東京都港区芝5丁目26-16　Mita S-Garden 2階",
        "　株式会社エール",
        "　代表取締役　矢島　英一　　　　　　　㊞",
    ]:
        add_paragraph_with_style(doc, line, size=11, space_after=2)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 乙
    add_paragraph_with_style(doc, "乙", bold=True, size=11, space_before=4, space_after=3)
    for line in [
        "　神奈川県横浜市港北区富士塚2-13-28",
        "　株式会社TOKUMORI",
        "　代表取締役　岡見　悠平　　　　　　　㊞",
    ]:
        add_paragraph_with_style(doc, line, size=11, space_after=2)

    # 保存
    output_path = "/Users/atsuyasato130/Claude AI/覚書_人材紹介手数料変更_2027年度新卒採用.docx"
    doc.save(output_path)
    print(f"保存完了: {output_path}")


if __name__ == "__main__":
    create_memorandum()
