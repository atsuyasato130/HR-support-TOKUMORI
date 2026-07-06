"""新卒採用_人材紹介手数料に関する覚書（27卒適用）を作成."""

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_jp_font(run, size_pt=10.5, bold=False):
    run.font.name = "MS Mincho"
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        from docx.oxml import OxmlElement
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "MS Mincho")
    r_fonts.set(qn("w:ascii"), "MS Mincho")
    r_fonts.set(qn("w:hAnsi"), "MS Mincho")


def add_para(doc, text, size=10.5, bold=False, align=None, indent_cm=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_jp_font(run, size, bold)
    return p


def add_heading(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_jp_font(run, size, bold=True)
    return p


def add_uni_table(doc, rows_data):
    """rows_data: list of rows. Each row is list of cells (up to 3 cols)."""
    table = doc.add_table(rows=len(rows_data), cols=3)
    table.style = "Table Grid"
    table.autofit = False
    for i, row_data in enumerate(rows_data):
        cells = table.rows[i].cells
        for j in range(3):
            cell = cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = row_data[j] if j < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_jp_font(run, 10)
    return table


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # タイトル
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("新卒採用_人材紹介手数料に関する覚書（27卒適用）")
    set_jp_font(run, 14, bold=True)

    # 前文
    add_para(
        doc,
        "株式会社ファミリーコーポレーション（以下「甲」という）と株式会社TOKUMORI（以下「乙」という）"
        "とは、甲乙間で2025年6月27日付にて締結した「人材紹介に関する基本契約書」（以下「原契約」という）"
        "第7条第4項に基づき、紹介手数料の単価に関し、以下のとおり合意する（以下「本覚書」という）。",
    )
    add_para(doc, "")

    # 第1条
    add_heading(doc, "第1条（適用対象）")
    add_para(
        doc,
        "本覚書は、2027年3月に大学（学部・大学院を含む）を卒業見込みの者（以下「27卒」という）であって、"
        "乙の紹介により甲が正社員雇用での内定を決定し、当該内定承諾日が2026年5月1日以降である者"
        "（以下「本対象者」という）に限り適用する。",
        indent_cm=0.5,
    )

    # 第2条
    add_heading(doc, "第2条（紹介手数料の単価）")
    add_para(
        doc,
        "本対象者に係る紹介手数料は、原契約第7条第2項および第3項の定めにかかわらず、"
        "本対象者の最終学歴である大学に応じて、次の各号に定める金額（いずれも消費税及び地方消費税別）"
        "とする。",
        indent_cm=0.5,
    )
    add_para(doc, "")

    # 200万円
    add_para(doc, "(1) 一律 200万円（税抜）", bold=True, indent_cm=0.5)
    add_uni_table(
        doc,
        [
            ["《国立》", "", ""],
            ["東京大学", "京都大学", "北海道大学"],
            ["東北大学", "名古屋大学", "大阪大学"],
            ["九州大学", "一橋大学", "東京科学大学"],
            ["神戸大学", "国際教養大学", "お茶の水女子大学"],
            ["筑波大学", "千葉大学", "大阪公立大学"],
            ["東京外国語大学", "", ""],
            ["《私立》", "", ""],
            ["早稲田大学", "慶應義塾大学", "上智大学"],
            ["ICU（国際基督教大学）", "東京理科大学", ""],
        ],
    )
    add_para(doc, "")

    # 150万円
    add_para(doc, "(2) 一律 150万円（税抜）", bold=True, indent_cm=0.5)
    add_uni_table(
        doc,
        [
            ["《国立》", "", ""],
            ["東京都立大学", "横浜国立大学", "金沢大学"],
            ["岡山大学", "広島大学", "電気通信大学"],
            ["東京農工大学", "名古屋工業大学", "京都工芸繊維大学"],
            ["横浜市立大学", "", ""],
            ["《私立》", "", ""],
            ["学習院大学", "明治大学", "青山学院大学"],
            ["立教大学", "中央大学", "法政大学"],
            ["関西学院大学", "関西大学", "同志社大学"],
            ["立命館大学", "", ""],
        ],
    )
    add_para(doc, "")

    # 130万円
    add_para(doc, "(3) 上記(1)(2)以外の大学", bold=True, indent_cm=0.5)
    add_para(doc, "一律 130万円（税抜）", indent_cm=1.0)
    add_para(doc, "")

    # 第3条
    add_heading(doc, "第3条（適用範囲および原契約との関係）")
    add_para(
        doc,
        "1. 本覚書は、本対象者に係る正社員雇用での内定承諾に対する紹介手数料についてのみ適用する。"
        "原契約第7条第1項（非正規雇用・インターン等に係る20万円）および同条第3項に基づく追加支払いに"
        "ついては、本覚書の適用対象外とし、原契約の定めによる。",
        indent_cm=0.5,
    )
    add_para(
        doc,
        "2. 本覚書に定めのない事項は、すべて原契約の定めるところによる。本覚書と原契約の内容に相違が"
        "ある場合には、本対象者に関する事項に限り、本覚書の定めを優先して適用する。",
        indent_cm=0.5,
    )

    # 第4条
    add_heading(doc, "第4条（有効期間）")
    add_para(
        doc,
        "本覚書は2026年5月1日より効力を生じ、原契約の有効期間中、本対象者に対して適用するものとする。"
        "原契約が終了した場合、本覚書も同時に終了する。",
        indent_cm=0.5,
    )

    # 第5条
    add_heading(doc, "第5条（協議）")
    add_para(
        doc,
        "本覚書の解釈に疑義が生じた場合、または本覚書に定めのない事項について必要が生じた場合には、"
        "甲乙誠実に協議の上、これを解決するものとする。",
        indent_cm=0.5,
    )

    add_para(doc, "")
    add_para(
        doc,
        "本覚書の成立を証するため、本書の電磁的記録を作成し、甲及び乙が合意の後電子署名を施し、"
        "各自その電磁的記録を保管する。",
    )

    # 日付・署名欄
    add_para(doc, "")
    add_para(doc, "2026年5月1日", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, "")

    # 署名テーブル
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.autofit = False
    sig_table.columns[0].width = Cm(1.5)
    sig_table.columns[1].width = Cm(14.0)

    # 甲
    cell_label = sig_table.rows[0].cells[0]
    cell_label.text = ""
    p = cell_label.paragraphs[0]
    run = p.add_run("甲")
    set_jp_font(run, 11)

    cell_co = sig_table.rows[0].cells[1]
    cell_co.text = ""
    for line in ["東京都中央区銀座6-10-1", "株式会社ファミリーコーポレーション", "代表取締役  冨吉  範明"]:
        p = cell_co.add_paragraph() if cell_co.paragraphs[0].text or line != "東京都中央区銀座6-10-1" else cell_co.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        set_jp_font(run, 11)

    # 乙
    cell_label = sig_table.rows[1].cells[0]
    cell_label.text = ""
    p = cell_label.paragraphs[0]
    run = p.add_run("乙")
    set_jp_font(run, 11)

    cell_co = sig_table.rows[1].cells[1]
    cell_co.text = ""
    for line in ["神奈川県横浜市港北区富士塚2-13-28", "株式会社TOKUMORI", "代表取締役  岡見  悠平"]:
        p = cell_co.add_paragraph() if cell_co.paragraphs[0].text or line != "神奈川県横浜市港北区富士塚2-13-28" else cell_co.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        set_jp_font(run, 11)

    out_path = "/Users/atsuyasato130/Claude AI/新卒採用_人材紹介手数料に関する覚書_27卒適用_ファミリーコーポレーション様.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
